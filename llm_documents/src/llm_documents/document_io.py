from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
from dataclasses import dataclass
from datetime import datetime
from importlib.resources import files
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any


SUPPORTED_READ_EXTENSIONS = {".docx", ".xlsx", ".pdf", ".md"}


@dataclass(frozen=True)
class DocumentSnapshot:
    path: str
    kind: str
    content: str
    parser: str


@dataclass(frozen=True)
class PdfMarkdownCacheResult:
    source_path: str
    markdown_path: str
    metadata_path: str
    cache_hit: bool
    content: str


def read_document(path: Path, *, prefer_docling: bool = True) -> DocumentSnapshot:
    """Read a supported document into text or Markdown."""
    path = path.expanduser().resolve()
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_READ_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {path.suffix}")

    if suffix == ".md":
        return DocumentSnapshot(
            path=str(path),
            kind=suffix.lstrip("."),
            content=path.read_text(encoding="utf-8"),
            parser="markdown",
        )

    if suffix == ".pdf":
        return DocumentSnapshot(
            path=str(path),
            kind=suffix.lstrip("."),
            content=_read_pdf_with_pypdf(path),
            parser="pypdf",
        )

    if prefer_docling:
        try:
            return DocumentSnapshot(
                path=str(path),
                kind=suffix.lstrip("."),
                content=_read_with_docling(path),
                parser="docling",
            )
        except Exception as exc:
            fallback = _read_with_fallback(path)
            return DocumentSnapshot(
                path=str(path),
                kind=suffix.lstrip("."),
                content=f"{fallback}\n\n[Docling fallback reason: {exc}]",
                parser=f"fallback:{_fallback_name(suffix)}",
            )

    return DocumentSnapshot(
        path=str(path),
        kind=suffix.lstrip("."),
        content=_read_with_fallback(path),
        parser=f"fallback:{_fallback_name(suffix)}",
    )


def extract_pdf_to_markdown_cache(*, pdf_path: Path, cache_dir: Path) -> PdfMarkdownCacheResult:
    """Extract a PDF once and reuse a Markdown cache while the source is unchanged."""
    pdf_path = pdf_path.expanduser().resolve()
    if pdf_path.suffix.lower() != ".pdf":
        raise ValueError(f"Expected a PDF file, got: {pdf_path.suffix}")

    cache_dir = cache_dir.expanduser().resolve()
    cache_dir.mkdir(parents=True, exist_ok=True)

    cache_stem = _pdf_cache_stem(pdf_path)
    markdown_path = cache_dir / f"{cache_stem}.md"
    metadata_path = cache_dir / f"{cache_stem}.json"
    metadata = _pdf_cache_metadata(pdf_path)

    if markdown_path.exists() and metadata_path.exists():
        try:
            cache_is_current = json.loads(metadata_path.read_text(encoding="utf-8")) == metadata
        except json.JSONDecodeError:
            cache_is_current = False
        if cache_is_current:
            return PdfMarkdownCacheResult(
                source_path=str(pdf_path),
                markdown_path=str(markdown_path),
                metadata_path=str(metadata_path),
                cache_hit=True,
                content=markdown_path.read_text(encoding="utf-8"),
            )

    content = _read_pdf_with_pypdf(pdf_path)
    markdown_path.write_text(content, encoding="utf-8")
    metadata_path.write_text(json.dumps(metadata, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    return PdfMarkdownCacheResult(
        source_path=str(pdf_path),
        markdown_path=str(markdown_path),
        metadata_path=str(metadata_path),
        cache_hit=False,
        content=content,
    )


def render_docx_brief(
    *,
    output_path: Path,
    title: str,
    question: str,
    answer: str,
    sources: list[DocumentSnapshot],
    template_path: Path | None = None,
) -> Path:
    """Render an edited DOCX brief using docxtpl."""
    from docxtpl import DocxTemplate

    output_path = output_path.expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if template_path:
        resolved_template_path = template_path.expanduser().resolve()
        if not resolved_template_path.exists():
            create_brief_template(resolved_template_path)
        _render_docx_template(
            template_path=resolved_template_path,
            output_path=output_path,
            title=title,
            question=question,
            answer=answer,
            sources=sources,
        )
        return output_path

    with TemporaryDirectory(prefix="llm-documents-template-") as temp_dir:
        resolved_template_path = Path(temp_dir) / "brief_template.docx"
        create_brief_template(resolved_template_path)
        _render_docx_template(
            template_path=resolved_template_path,
            output_path=output_path,
            title=title,
            question=question,
            answer=answer,
            sources=sources,
        )
    return output_path


def _render_docx_template(
    *,
    template_path: Path,
    output_path: Path,
    title: str,
    question: str,
    answer: str,
    sources: list[DocumentSnapshot],
) -> None:
    from docxtpl import DocxTemplate

    template = DocxTemplate(str(template_path))
    template.render(
        {
            "title": title,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "question": question,
            "answer": answer,
            "sources": [
                {
                    "name": Path(source.path).name,
                    "kind": source.kind,
                    "parser": source.parser,
                    "chars": len(source.content),
                }
                for source in sources
            ],
        }
    )
    template.save(str(output_path))


def create_brief_template(path: Path) -> Path:
    """Create a minimal docxtpl-compatible DOCX template."""
    from docx import Document

    path = path.expanduser().resolve()
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("{{ title }}", level=1)
    document.add_paragraph("Generated: {{ generated_at }}")
    document.add_heading("Question", level=2)
    document.add_paragraph("{{ question }}")
    document.add_heading("Local LLM Synthesis", level=2)
    document.add_paragraph("{{ answer }}")
    document.add_heading("Sources", level=2)
    document.add_paragraph(
        "{% for source in sources %}"
        "{{ source.name }} | {{ source.kind }} | {{ source.parser }} | {{ source.chars }} chars\n"
        "{% endfor %}"
    )
    document.save(path)
    return path


def write_xlsx_report(
    *,
    output_path: Path,
    question: str,
    answer: str,
    sources: list[DocumentSnapshot],
    base_workbook_path: Path | None = None,
) -> Path:
    """Create or edit an XLSX workbook containing the LLM result and source metrics."""
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    output_path = output_path.expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if base_workbook_path and base_workbook_path.exists():
        workbook = load_workbook(base_workbook_path)
        if "LLM Summary" in workbook.sheetnames:
            del workbook["LLM Summary"]
        ws = workbook.create_sheet("LLM Summary", 0)
    else:
        workbook = Workbook()
        ws = workbook.active
        ws.title = "LLM Summary"

    ws["A1"] = "Question"
    ws["B1"] = question
    ws["A2"] = "Answer"
    ws["B2"] = answer
    ws["A4"] = "Source"
    ws["B4"] = "Kind"
    ws["C4"] = "Parser"
    ws["D4"] = "Characters"

    for row_idx, source in enumerate(sources, start=5):
        ws.cell(row=row_idx, column=1, value=Path(source.path).name)
        ws.cell(row=row_idx, column=2, value=source.kind)
        ws.cell(row=row_idx, column=3, value=source.parser)
        ws.cell(row=row_idx, column=4, value=len(source.content))

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for row in (1, 2, 4):
        for cell in ws[row]:
            if cell.value:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(vertical="top")

    ws["B2"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 72
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 14
    ws.row_dimensions[2].height = 160

    workbook.save(output_path)
    return output_path


def export_to_pdf(*, input_path: Path, output_dir: Path, timeout_seconds: int = 60) -> Path:
    """Export an office document to PDF using local LibreOffice."""
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if not executable:
        raise RuntimeError("LibreOffice was not found on PATH.")

    input_path = input_path.expanduser().resolve()
    output_dir = output_dir.expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{input_path.stem}.pdf"

    with TemporaryDirectory(prefix="llm-documents-libreoffice-") as profile_dir:
        try:
            result = subprocess.run(
                [
                    executable,
                    f"-env:UserInstallation=file://{profile_dir}",
                    "--headless",
                    "--nologo",
                    "--nofirststartwizard",
                    "--nolockcheck",
                    "--nodefault",
                    "--convert-to",
                    _libreoffice_pdf_filter(input_path),
                    "--outdir",
                    str(output_dir),
                    str(input_path),
                ],
                capture_output=True,
                check=False,
                text=True,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(f"LibreOffice PDF export timed out after {timeout_seconds}s") from exc

    if result.returncode != 0:
        details = (result.stderr or result.stdout).strip()
        raise RuntimeError(f"LibreOffice PDF export failed with exit code {result.returncode}: {details}")
    if not output_path.exists():
        raise RuntimeError(f"LibreOffice did not create expected output: {output_path}")
    return output_path


def _libreoffice_pdf_filter(input_path: Path) -> str:
    suffix = input_path.suffix.lower()
    if suffix in {".doc", ".docx", ".odt"}:
        return "pdf:writer_pdf_Export"
    if suffix in {".xls", ".xlsx", ".ods"}:
        return "pdf:calc_pdf_Export"
    return "pdf"


def create_demo_documents(output_dir: Path) -> list[Path]:
    """Create sample DOCX, XLSX, and PDF inputs for the demo workflow."""
    output_dir = output_dir.expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = output_dir / "bios_support_notes.docx"
    xlsx_path = output_dir / "bios_triage_workbook.xlsx"
    pdf_path = output_dir / "asus_bios_manual.pdf"

    _create_demo_docx(docx_path)
    _create_demo_xlsx(xlsx_path)
    _copy_demo_pdf(pdf_path)

    return [docx_path, xlsx_path, pdf_path]


def _read_with_docling(path: Path) -> str:
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(str(path))
    return result.document.export_to_markdown()


def _read_with_fallback(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx_with_python_docx(path)
    if suffix == ".xlsx":
        return _read_xlsx_with_openpyxl(path)
    if suffix == ".pdf":
        return _read_pdf_with_pypdf(path)
    raise ValueError(f"No fallback reader for {path.suffix}")


def _fallback_name(suffix: str) -> str:
    return {
        ".docx": "python-docx",
        ".xlsx": "openpyxl",
        ".pdf": "pypdf",
    }[suffix]


def _read_docx_with_python_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_idx, table in enumerate(document.tables, start=1):
        parts.append(f"\nTable {table_idx}")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))

    return "\n".join(parts)


def _read_xlsx_with_openpyxl(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for worksheet in workbook.worksheets:
        parts.append(f"# Sheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                parts.append(" | ".join(values))
        parts.append("")
    return "\n".join(parts).strip()


def _read_pdf_with_pypdf(path: Path) -> str:
    from pypdf import PdfReader

    parts: list[str] = []
    reader = PdfReader(path)
    for page_idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"# Page {page_idx}\n{text}")
    return "\n\n".join(parts)


def _pdf_cache_stem(path: Path) -> str:
    digest = hashlib.sha256(str(path).encode("utf-8")).hexdigest()[:10]
    return f"{path.stem}-{digest}"


def _pdf_cache_metadata(path: Path) -> dict[str, Any]:
    stat = path.stat()
    return {
        "format_version": 1,
        "parser": "pypdf",
        "source_path": str(path),
        "source_size": stat.st_size,
        "source_mtime_ns": stat.st_mtime_ns,
    }


def _create_demo_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("BIOS Support Escalation Notes", level=1)
    document.add_paragraph(
        "A workstation fleet using ASUS PRIME / ProArt / TUF GAMING AMD AM5 boards "
        "shows intermittent boot instability after manual BIOS tuning."
    )
    document.add_paragraph(
        "Support needs a local-only summary of safe recovery steps, fan-control checks, "
        "and which BIOS areas require engineering approval before changes."
    )
    document.add_paragraph("Target response date: 2026-06-12.")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ["Owner", "Item", "Status"], strict=True):
        cell.text = header
    rows = [
        ("Mina", "Collect affected BIOS versions", "In progress"),
        ("Jon", "Document optimized-defaults recovery path", "At risk"),
        ("Lea", "Review QFan and boot-setting guidance", "In progress"),
    ]
    for owner, item, status in rows:
        cells = table.add_row().cells
        cells[0].text = owner
        cells[1].text = item
        cells[2].text = status
    document.save(path)


def _create_demo_xlsx(path: Path) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    ws = workbook.active
    ws.title = "Triage"
    ws.append(["Metric", "Current", "Target", "Comment"])
    rows: list[tuple[Any, ...]] = [
        ("Affected workstations reviewed", 18, 40, "Need more inventory coverage"),
        ("Open BIOS incidents", 17, 5, "Needs triage"),
        ("Machines recovered with defaults", 0.68, 0.9, "Below target"),
    ]
    for row in rows:
        ws.append(row)

    risks = workbook.create_sheet("Risks")
    risks.append(["Risk", "Impact", "Mitigation"])
    risks.append(["Unsafe BIOS tuning", "High", "Require engineering approval for voltage or timing changes"])
    risks.append(["Unclear recovery process", "High", "Document optimized defaults and CMOS reset escalation path"])
    risks.append(["Fan profile mismatch", "Medium", "Verify QFan profile, PWM/DC mode, and manual fan settings"])
    workbook.save(path)


def _copy_demo_pdf(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    source = files("llm_documents").joinpath(
        "assets/PRIME_PROART_TUF_GAMING_AMD_AM5_Series_BIOS_EM_WEB_DE.pdf"
    )
    with source.open("rb") as source_file, path.open("wb") as target_file:
        shutil.copyfileobj(source_file, target_file)
